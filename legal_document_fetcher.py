"""
Legal Document Fetcher Module

This module provides Object-Oriented classes for fetching Brazilian legal documents
from URLs and saving them as Microsoft Word documents.

Main classes:
- FetcherConfig: Configuration settings
- LegalDocumentFetcher: Main orchestrator class
- HTMLContentExtractor: Extract and clean HTML content
- WordDocumentBuilder: Convert HTML to Word documents
- FetchResult: Store results of fetch operations
"""

import logging
import os
import re
import time
import base64
import io
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional, Dict, Any
from urllib.parse import urlparse, parse_qs

import requests
from bs4 import BeautifulSoup, Comment
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from tqdm import tqdm
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from webdriver_manager.chrome import ChromeDriverManager


# Configure module logger
logger = logging.getLogger(__name__)
logger.setLevel(logging.INFO)

# Create console handler if not already present
if not logger.handlers:
    handler = logging.StreamHandler()
    handler.setLevel(logging.INFO)
    formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
    handler.setFormatter(formatter)
    logger.addHandler(handler)


@dataclass
class FetcherConfig:
    """Configuration settings for the Legal Document Fetcher."""

    output_dir: str = "./legal_documents"
    request_timeout: int = 30
    retry_attempts: int = 3
    delay_between_requests: float = 2.0
    user_agent: str = "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
    content_selector: str = "div.texto"
    create_output_dir: bool = True
    use_selenium: bool = True  # Use Selenium for JavaScript-rendered pages
    selenium_wait_time: int = 20  # Max wait time for page elements to load (increased for reliability)

    def __post_init__(self):
        """Create output directory if it doesn't exist."""
        if self.create_output_dir:
            Path(self.output_dir).mkdir(parents=True, exist_ok=True)


@dataclass
class FetchResult:
    """Store the result of a single fetch operation."""

    url: str
    success: bool
    law_number: str = ""
    filename: str = ""
    error_message: Optional[str] = None
    fetch_time: float = 0.0

    def __str__(self):
        if self.success:
            return f"✓ {self.law_number} -> {self.filename} ({self.fetch_time:.2f}s)"
        else:
            return f"✗ {self.url} - Error: {self.error_message}"


class HTMLContentExtractor:
    """Extract and clean HTML content from legal document webpages."""

    def __init__(self, content_selector: str = "div.texto"):
        """
        Initialize the HTML content extractor.

        Args:
            content_selector: CSS selector for the main content div
        """
        self.content_selector = content_selector
        self.fallback_selectors = [
            "app-legislacao",  # Angular component for normas.leg.br
            "div.texto",
            "div#texto",
            "article",
            "main",
            "div.content",
            "div#content",
            "div.container"
        ]

    def extract_main_content(self, html: str) -> Optional[BeautifulSoup]:
        """
        Extract the main content from HTML using CSS selectors.
        Specifically handles Shadow DOM content markers.

        Args:
            html: Raw HTML string

        Returns:
            BeautifulSoup object with main content or None if not found
        """
        try:
            soup = BeautifulSoup(html, 'html.parser')

            # Check if this HTML contains shadow DOM content
            if '<!-- SHADOW DOM CONTENT -->' in html:
                logger.info("Detected shadow DOM content marker")
                # Extract just the shadow DOM part
                shadow_start = html.find('<!-- SHADOW DOM CONTENT -->')
                shadow_end = html.find('<!-- END SHADOW DOM -->')

                if shadow_start != -1 and shadow_end != -1:
                    shadow_html = html[shadow_start:shadow_end]
                    shadow_soup = BeautifulSoup(shadow_html, 'html.parser')

                    # Remove style tags (we don't need CSS in the document)
                    for style in shadow_soup.find_all('style'):
                        style.decompose()

                    # Find the main content container div (usually has classes like 'p-2', 'w-100', 'border-0')
                    # This is typically the only non-style element in the shadow root
                    main_div = None
                    for elem in shadow_soup.children:
                        if hasattr(elem, 'name') and elem.name == 'div':
                            main_div = elem
                            break

                    if main_div:
                        logger.info(f"Found main content div in shadow DOM")
                        # Return the main div which contains all the actual content
                        return main_div
                    else:
                        logger.warning("Could not find main content div in shadow DOM")
                        # Fallback: return the shadow soup itself
                        return shadow_soup

            # Try primary selector
            content = soup.select_one(self.content_selector)

            # Try fallback selectors if primary fails
            if not content:
                logger.warning(f"Primary selector '{self.content_selector}' failed, trying fallbacks")
                for selector in self.fallback_selectors:
                    content = soup.select_one(selector)
                    if content:
                        logger.info(f"Found content with fallback selector: {selector}")
                        break

            # Last resort: return body content
            if not content:
                logger.warning("All selectors failed, using body content")
                content = soup.body if soup.body else soup

            return content

        except Exception as e:
            logger.error(f"Error parsing HTML: {e}")
            return None

    def clean_content(self, soup: BeautifulSoup) -> BeautifulSoup:
        """
        Clean up HTML content by removing unnecessary elements.

        Args:
            soup: BeautifulSoup object to clean

        Returns:
            Cleaned BeautifulSoup object
        """
        # Remove script and style elements
        for element in soup(['script', 'style', 'meta', 'link', 'noscript']):
            element.decompose()

        # Remove comments
        for comment in soup.find_all(string=lambda text: isinstance(text, Comment)):
            comment.extract()

        # Remove empty tags
        for tag in soup.find_all():
            if not tag.get_text(strip=True) and not tag.name in ['br', 'hr', 'img']:
                tag.decompose()

        return soup

    def get_law_title(self, soup: BeautifulSoup) -> str:
        """
        Extract the law title from content.

        Args:
            soup: BeautifulSoup object with content

        Returns:
            Law title string or generic title
        """
        # Try to find title in common heading tags
        for tag in ['h1', 'h2', 'title']:
            title_elem = soup.find(tag)
            if title_elem:
                title = title_elem.get_text(strip=True)
                if title:
                    return title

        # Fallback: look for text that looks like a law title
        text = soup.get_text()
        match = re.search(r'Lei[^0-9]+\d+[^\n]+', text)
        if match:
            return match.group(0).strip()

        return "Legal Document"


class WordDocumentBuilder:
    """Convert HTML content to Microsoft Word documents."""

    def __init__(self):
        """Initialize the Word document builder."""
        pass

    def create_document(self, content: BeautifulSoup, title: str) -> Document:
        """
        Create a new Word document with the given content.

        Args:
            content: BeautifulSoup object with HTML content
            title: Document title

        Returns:
            Document object
        """
        doc = Document()

        # Set document properties (title has 255 char limit in Word)
        truncated_title = title[:255] if len(title) > 255 else title
        doc.core_properties.title = truncated_title
        doc.core_properties.subject = "Brazilian Federal Law"

        # Only add title heading if it's meaningful (not generic or too long)
        # Skip title if it's the generic "Legal Document" or very long (likely extracted incorrectly)
        if title != "Legal Document" and len(title) < 200:
            logger.info(f"Adding document title: {title[:100]}...")
            title_para = doc.add_heading(title, level=1)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            logger.info(f"Skipping title heading (generic or too long): {title[:100]}...")

        # Add content
        logger.info("Starting to add HTML content to document")
        self.add_html_content(doc, content)
        logger.info(f"Finished adding content. Document now has {len(doc.paragraphs)} paragraphs")

        return doc

    def add_html_content(self, doc: Document, soup: BeautifulSoup) -> None:
        """
        Convert HTML elements to Word formatting and add to document.

        Args:
            doc: Document object to add content to
            soup: BeautifulSoup object with HTML content
        """
        # Process each top-level element
        for element in soup.children:
            if element.name is None:
                # Text node
                text = element.strip()
                if text:
                    doc.add_paragraph(text)

            elif element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                # Headings
                level = int(element.name[1])
                text = element.get_text(strip=True)
                if text:
                    doc.add_heading(text, level=min(level, 9))

            elif element.name == 'p':
                # Paragraphs
                # Check if paragraph contains an image
                img = element.find('img')
                if img:
                    logger.info("Found image in paragraph, adding to document")
                    # Add the image first
                    self._add_image(doc, img)

                # Then add the text content (if any)
                text = element.get_text(strip=True)
                if text:
                    para = doc.add_paragraph()
                    self._add_formatted_text(para, element)

            elif element.name in ['ul', 'ol']:
                # Lists
                self._add_list(doc, element)

            elif element.name == 'table':
                # Tables
                self._add_table(doc, element)

            elif element.name == 'div':
                # Recursively process div contents
                self.add_html_content(doc, element)

            elif element.name == 'img':
                # Image
                self._add_image(doc, element)

            elif element.name == 'br':
                # Line break
                doc.add_paragraph()

            else:
                # Default: add as paragraph
                text = element.get_text(strip=True)
                if text and len(text) > 1:
                    doc.add_paragraph(text)

    def _add_formatted_text(self, para, element):
        """Add text with formatting (bold, italic) to paragraph."""
        for content in element.children:
            if content.name is None:
                # Plain text
                text = str(content)
                if text.strip():
                    para.add_run(text)

            elif content.name == 'strong' or content.name == 'b':
                # Bold
                run = para.add_run(content.get_text())
                run.bold = True

            elif content.name == 'em' or content.name == 'i':
                # Italic
                run = para.add_run(content.get_text())
                run.italic = True

            elif content.name == 'u':
                # Underline
                run = para.add_run(content.get_text())
                run.underline = True

            elif content.name == 'img':
                # Image inside paragraph - need to handle specially
                # Images cannot be added inline to an existing paragraph with text
                # So we need to handle this in the parent method
                pass

            else:
                # Other tags - just add text
                text = content.get_text()
                if text.strip():
                    para.add_run(text)

    def _add_list(self, doc: Document, list_element):
        """Add a list (ul or ol) to the document."""
        for li in list_element.find_all('li', recursive=False):
            text = li.get_text(strip=True)
            if text:
                doc.add_paragraph(text, style='List Bullet' if list_element.name == 'ul' else 'List Number')

    def _add_image(self, doc: Document, img_element):
        """Add an image to the document from base64 data or URL."""
        try:
            src = img_element.get('src', '')

            if not src:
                return

            # Handle base64-encoded images
            if src.startswith('data:image'):
                # Extract base64 data
                # Format: data:image/png;base64,iVBORw0KG...
                if ';base64,' in src:
                    base64_data = src.split(';base64,')[1]
                    image_data = base64.b64decode(base64_data)
                    image_stream = io.BytesIO(image_data)

                    # Add image to document (width in inches, 68px ~= 0.71 inches at 96 DPI)
                    para = doc.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.add_run()
                    run.add_picture(image_stream, width=Inches(0.71))
                    logger.info("Added base64 image to document")
                else:
                    logger.warning(f"Unsupported data URI format: {src[:50]}...")
            else:
                # TODO: Handle regular URLs if needed
                logger.info(f"Skipping external image URL: {src}")

        except Exception as e:
            logger.warning(f"Failed to add image: {e}")

    def _add_table(self, doc: Document, table_element):
        """Add a table to the document."""
        rows = table_element.find_all('tr')
        if not rows:
            return

        # Count columns
        max_cols = max(len(row.find_all(['td', 'th'])) for row in rows)

        # Create table
        table = doc.add_table(rows=len(rows), cols=max_cols)
        table.style = 'Light Grid Accent 1'

        # Fill table
        for i, row in enumerate(rows):
            cells = row.find_all(['td', 'th'])
            for j, cell in enumerate(cells):
                if j < max_cols:
                    table.rows[i].cells[j].text = cell.get_text(strip=True)

    def save_document(self, doc: Document, filepath: str) -> None:
        """
        Save the document to a file.

        Args:
            doc: Document object to save
            filepath: Path where to save the file

        Raises:
            IOError: If unable to save the file
        """
        try:
            doc.save(filepath)
            logger.info(f"Saved document: {filepath}")
        except Exception as e:
            logger.error(f"Error saving document to {filepath}: {e}")
            raise IOError(f"Failed to save document: {e}")


class LegalDocumentFetcher:
    """Main class to orchestrate fetching and saving legal documents."""

    def __init__(self, config: FetcherConfig = None):
        """
        Initialize the Legal Document Fetcher.

        Args:
            config: FetcherConfig object with settings
        """
        self.config = config if config else FetcherConfig()

        # Initialize session with headers
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': self.config.user_agent,
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
            'Accept-Language': 'pt-BR,pt;q=0.9,en;q=0.8',
        })

        # Initialize helper objects
        self.content_extractor = HTMLContentExtractor(self.config.content_selector)
        self.doc_builder = WordDocumentBuilder()

        # Store results
        self.results: List[FetchResult] = []

        # Selenium driver (lazy initialization)
        self.driver = None

    def _init_selenium_driver(self):
        """Initialize Selenium WebDriver with Chrome in headless mode."""
        if self.driver is None:
            try:
                chrome_options = Options()
                chrome_options.add_argument('--headless=new')  # Use new headless mode for better compatibility
                chrome_options.add_argument('--no-sandbox')
                chrome_options.add_argument('--disable-dev-shm-usage')
                chrome_options.add_argument('--disable-gpu')
                chrome_options.add_argument('--disable-blink-features=AutomationControlled')
                chrome_options.add_argument('--window-size=1920,1080')
                chrome_options.add_argument(f'user-agent={self.config.user_agent}')

                # Use webdriver-manager to automatically handle driver installation
                service = Service(ChromeDriverManager().install())
                self.driver = webdriver.Chrome(service=service, options=chrome_options)
                logger.info("Selenium WebDriver initialized successfully")
            except Exception as e:
                logger.error(f"Failed to initialize Selenium driver: {e}")
                raise

    def _fetch_with_selenium(self, url: str) -> Optional[str]:
        """
        Fetch webpage using Selenium for JavaScript-rendered content.
        Handles Shadow DOM elements.

        Args:
            url: URL to fetch

        Returns:
            HTML string with rendered content or None if failed
        """
        try:
            self._init_selenium_driver()

            logger.info(f"Fetching with Selenium: {url}")
            self.driver.get(url)

            # Try to access Shadow DOM content if present
            html = self._get_html_with_shadow_dom()

            logger.info(f"Successfully fetched {len(html)} bytes with Selenium")
            return html

        except Exception as e:
            logger.error(f"Error fetching with Selenium: {e}")
            return None

    def _get_html_with_shadow_dom(self) -> str:
        """
        Get HTML including content from Shadow DOM elements.
        Waits for shadow DOM elements to be present before extracting.

        Returns:
            HTML string with shadow DOM content expanded
        """
        try:
            # Wait for the shadow host element to be present
            wait = WebDriverWait(self.driver, self.config.selenium_wait_time)

            try:
                # Wait for sf-unstructured-legislation-viewer element
                element = wait.until(
                    EC.presence_of_element_located((By.TAG_NAME, "sf-unstructured-legislation-viewer"))
                )
                logger.info("Found sf-unstructured-legislation-viewer element")

                # Access the shadow root
                shadow_root = self.driver.execute_script('return arguments[0].shadowRoot', element)

                if shadow_root:
                    # Get the innerHTML of the shadow root
                    shadow_html = self.driver.execute_script('return arguments[0].innerHTML', shadow_root)

                    if shadow_html:
                        logger.info(f"Extracted {len(shadow_html)} bytes from shadow DOM")

                        # Get base page source
                        base_html = self.driver.page_source

                        # Append shadow DOM content to the base HTML
                        return base_html + f"\n<!-- SHADOW DOM CONTENT -->\n{shadow_html}\n<!-- END SHADOW DOM -->\n"
                    else:
                        logger.warning("Shadow root was empty")
                else:
                    logger.warning("Could not access shadow root")

            except Exception as e:
                logger.info(f"Shadow DOM element not found or timeout: {e}")
                logger.info("Returning page source without shadow DOM content")

            # Return base HTML if shadow DOM extraction failed
            return self.driver.page_source

        except Exception as e:
            logger.warning(f"Error extracting shadow DOM content: {e}")
            return self.driver.page_source

    def fetch_webpage(self, url: str) -> Optional[str]:
        """
        Fetch HTML content from URL with retry logic.
        Uses Selenium for JavaScript-rendered pages if configured.

        Args:
            url: URL to fetch

        Returns:
            HTML string or None if failed
        """
        # Use Selenium if configured
        if self.config.use_selenium:
            return self._fetch_with_selenium(url)

        # Otherwise use traditional requests
        for attempt in range(1, self.config.retry_attempts + 1):
            try:
                logger.info(f"Fetching {url} (attempt {attempt}/{self.config.retry_attempts})")
                response = self.session.get(url, timeout=self.config.request_timeout)
                response.raise_for_status()
                return response.text

            except requests.exceptions.Timeout:
                logger.warning(f"Timeout fetching {url} (attempt {attempt})")
                if attempt < self.config.retry_attempts:
                    time.sleep(2 ** attempt)  # Exponential backoff

            except requests.exceptions.HTTPError as e:
                logger.error(f"HTTP error fetching {url}: {e}")
                return None

            except requests.exceptions.ConnectionError as e:
                logger.warning(f"Connection error fetching {url} (attempt {attempt}): {e}")
                if attempt < self.config.retry_attempts:
                    time.sleep(2 ** attempt)

            except Exception as e:
                logger.error(f"Unexpected error fetching {url}: {e}")
                return None

        logger.error(f"Failed to fetch {url} after {self.config.retry_attempts} attempts")
        return None

    def extract_law_number_from_url(self, url: str) -> str:
        """
        Extract law number from LexML URN URL.

        Args:
            url: LexML URL

        Returns:
            Law number string (e.g., "lei_10101_20001219" for law 10101 from 2000-12-19)
        """
        try:
            # Parse URN from URL
            # Format: https://normas.leg.br/?urn=urn:lex:br:federal:lei:YYYY-MM-DD;NUMBER
            parsed = urlparse(url)
            urn = parse_qs(parsed.query).get('urn', [''])[0]

            if urn:
                # Extract number from URN (last part after semicolon)
                parts = urn.split(';')
                if len(parts) >= 2:
                    law_number = parts[-1]

                    # Extract full date from date part (YYYY-MM-DD)
                    date_part = parts[0].split(':')[-1]  # YYYY-MM-DD
                    # Remove hyphens to get YYYYMMDD format
                    date_formatted = date_part.replace('-', '')

                    return f"lei_{law_number}_{date_formatted}"

            # Fallback: use hash of URL
            return f"lei_{abs(hash(url)) % 100000}"

        except Exception as e:
            logger.warning(f"Error extracting law number from {url}: {e}")
            return f"lei_{abs(hash(url)) % 100000}"

    def generate_filename(self, law_number: str, url: str) -> str:
        """
        Generate a filename for the Word document.

        Args:
            law_number: Law number identifier
            url: Source URL

        Returns:
            Full filepath for the document
        """
        # Sanitize filename
        safe_name = re.sub(r'[^\w\-_]', '_', law_number)
        filename = f"{safe_name}.docx"

        filepath = os.path.join(self.config.output_dir, filename)

        # Handle duplicates by adding counter
        counter = 1
        while os.path.exists(filepath):
            filename = f"{safe_name}_{counter}.docx"
            filepath = os.path.join(self.config.output_dir, filename)
            counter += 1

        return filepath

    def process_single_url(self, url: str) -> FetchResult:
        """
        Process a single URL through the complete pipeline.

        Args:
            url: URL to process

        Returns:
            FetchResult object with operation details
        """
        start_time = time.time()

        try:
            # Extract law number
            law_number = self.extract_law_number_from_url(url)

            # Fetch webpage
            html = self.fetch_webpage(url)
            if not html:
                return FetchResult(
                    url=url,
                    success=False,
                    law_number=law_number,
                    error_message="Failed to fetch webpage",
                    fetch_time=time.time() - start_time
                )

            # Extract content
            content = self.content_extractor.extract_main_content(html)
            if not content:
                return FetchResult(
                    url=url,
                    success=False,
                    law_number=law_number,
                    error_message="Failed to extract content",
                    fetch_time=time.time() - start_time
                )

            # Clean content
            content = self.content_extractor.clean_content(content)

            # Check if we have meaningful content (not just empty page)
            text_content = content.get_text(strip=True)
            if len(text_content) < 100:
                logger.warning(f"Extracted content is too short ({len(text_content)} chars), likely failed to load properly")
                return FetchResult(
                    url=url,
                    success=False,
                    law_number=law_number,
                    error_message="Insufficient content extracted (shadow DOM may not have loaded)",
                    fetch_time=time.time() - start_time
                )

            # Get title
            title = self.content_extractor.get_law_title(content)

            # Create Word document
            doc = self.doc_builder.create_document(content, title)

            # Generate filename and save
            filepath = self.generate_filename(law_number, url)
            self.doc_builder.save_document(doc, filepath)

            # Success
            return FetchResult(
                url=url,
                success=True,
                law_number=law_number,
                filename=os.path.basename(filepath),
                fetch_time=time.time() - start_time
            )

        except Exception as e:
            logger.error(f"Error processing {url}: {e}")
            return FetchResult(
                url=url,
                success=False,
                law_number=law_number if 'law_number' in locals() else "unknown",
                error_message=str(e),
                fetch_time=time.time() - start_time
            )

    def process_url_list(self, urls: List[str], show_progress: bool = True) -> List[FetchResult]:
        """
        Process a list of URLs with progress tracking.

        Args:
            urls: List of URLs to process
            show_progress: Whether to show progress bar

        Returns:
            List of FetchResult objects
        """
        self.results = []

        try:
            iterator = tqdm(urls, desc="Fetching documents") if show_progress else urls

            for i, url in enumerate(iterator):
                result = self.process_single_url(url)
                self.results.append(result)

                # Log result
                if result.success:
                    logger.info(f"✓ {result.law_number} -> {result.filename}")
                else:
                    logger.error(f"✗ {url} - {result.error_message}")

                # Delay between requests (except for last one)
                if i < len(urls) - 1:
                    time.sleep(self.config.delay_between_requests)

        finally:
            # Clean up Selenium driver if it was used
            self.cleanup()

        return self.results

    def cleanup(self):
        """Clean up resources, especially Selenium driver."""
        if self.driver:
            try:
                self.driver.quit()
                logger.info("Selenium driver closed")
            except Exception as e:
                logger.warning(f"Error closing Selenium driver: {e}")
            finally:
                self.driver = None

    def get_summary(self) -> Dict[str, Any]:
        """
        Get summary statistics of the fetch operation.

        Returns:
            Dictionary with summary statistics
        """
        total = len(self.results)
        success = sum(1 for r in self.results if r.success)
        failed = total - success

        failed_urls = [r.url for r in self.results if not r.success]

        avg_time = sum(r.fetch_time for r in self.results) / total if total > 0 else 0

        return {
            'total': total,
            'success': success,
            'failed': failed,
            'success_rate': (success / total * 100) if total > 0 else 0,
            'failed_urls': failed_urls,
            'avg_fetch_time': avg_time
        }

    def export_results_to_csv(self, filepath: str) -> None:
        """
        Export results to a CSV file.

        Args:
            filepath: Path to save CSV file
        """
        try:
            import csv

            with open(filepath, 'w', newline='', encoding='utf-8') as f:
                writer = csv.writer(f)
                writer.writerow(['URL', 'Success', 'Law Number', 'Filename', 'Error', 'Fetch Time'])

                for result in self.results:
                    writer.writerow([
                        result.url,
                        result.success,
                        result.law_number,
                        result.filename,
                        result.error_message or '',
                        f"{result.fetch_time:.2f}"
                    ])

            logger.info(f"Results exported to {filepath}")

        except Exception as e:
            logger.error(f"Error exporting results to CSV: {e}")
            raise


# Example usage
if __name__ == "__main__":
    # Configuration
    config = FetcherConfig(
        output_dir="./legal_documents",
        delay_between_requests=2.0,
        retry_attempts=3
    )

    # Create fetcher
    fetcher = LegalDocumentFetcher(config)

    # Example URL
    test_urls = [
        "https://normas.leg.br/?urn=urn:lex:br:federal:lei:2000-12-19;10101"
    ]

    # Process
    results = fetcher.process_url_list(test_urls)

    # Summary
    summary = fetcher.get_summary()
    print(f"\nSummary:")
    print(f"Total: {summary['total']}")
    print(f"Success: {summary['success']}")
    print(f"Failed: {summary['failed']}")
    print(f"Success Rate: {summary['success_rate']:.2f}%")
